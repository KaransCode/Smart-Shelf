import sys
import os
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx is not installed. Please install it using pip.")
    sys.exit(1)

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('Smart-Shelf: Official Competition Technical Strategy & Synopsis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_heading('Executive Context', level=1)
    doc.add_paragraph(
        "Acting as Principal Solutions Architect and Official Judge for the National AI Impact Festival. "
        "The following document analyzes the 'Smart-Shelf: Consumer Food Waste & Expiry Predictor' project "
        "against the National Evaluation Rubrics (Metrics 01, 02, and 03) and provides precise technical "
        "architectures, code structures, and deployment strategies to secure maximum evaluation points."
    )

    # Deliverable 1
    doc.add_heading('Deliverable 1: Technical Fixes for Metric 01 (Inclusion & Accessibility)', level=1)
    
    doc.add_heading('1.1 WCAG Accessibility & Voice Assistance in React (TypeScript)', level=2)
    doc.add_paragraph(
        "To achieve maximum points in Inclusion, the UI must remove barriers for visually and motor-impaired users. "
        "We implement ARIA live regions to announce AI-generated recipes via screen readers automatically, and "
        "integrate the Web Speech API for hands-free camera toggling."
    )
    code_block_1 = doc.add_paragraph()
    code_block_1.add_run(
        "// React Code Snippet for Accessibility & Voice\n"
        "useEffect(() => {\n"
        "  const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;\n"
        "  if (SpeechRecognition) {\n"
        "    const recognition = new SpeechRecognition();\n"
        "    recognition.continuous = true;\n"
        "    recognition.onresult = (event) => {\n"
        "      const transcript = event.results[event.results.length - 1][0].transcript.trim().toLowerCase();\n"
        "      if (transcript.includes('start camera')) toggleCamera('start');\n"
        "      if (transcript.includes('stop camera')) toggleCamera('stop');\n"
        "    };\n"
        "    recognition.start();\n"
        "  }\n"
        "}, []);\n\n"
        "return (\n"
        "  <div aria-live=\"polite\" className=\"recipe-announcer\">\n"
        "    <span className=\"sr-only\">{newRecipeAnnouncment}</span>\n"
        "  </div>\n"
        ");"
    )
    code_block_1.style = 'Intense Quote'

    doc.add_heading('1.2 Offline Resiliency Strategy', level=2)
    doc.add_paragraph(
        "The Ultralytics YOLOv8 inference and decay regression algorithms run entirely locally on the Edge CPU/NPU. "
        "In offline/low-bandwidth scenarios, the Python script continues detecting items and communicating with "
        "the local Node.js MQTT broker. We implement a local fallback where, if the Gemini API cannot be reached, "
        "Node.js pulls from a locally cached dictionary of standard zero-waste recipes, ensuring zero downtime."
    )

    # Deliverable 2
    doc.add_heading('Deliverable 2: Architecture for Metric 02 (Responsible AI & Privacy)', level=1)
    
    doc.add_heading('2.1 Data Flattening Pipeline (100% Privacy Guarantee)', level=2)
    doc.add_paragraph(
        "To address Ethical Concerns and Data Privacy, we guarantee that raw camera frames never leave the device. "
        "The frame array is kept in volatile RAM, processed by YOLOv8, flattened into mathematical metrics, and immediately destroyed."
    )
    code_block_2 = doc.add_paragraph()
    code_block_2.add_run(
        "# Python Data Flattening Snippet\n"
        "ret, frame = cap.read()\n"
        "results = model(frame) # YOLOv8 Inference\n"
        "\n"
        "# FLATTENING: Extract only anonymous metrics\n"
        "telemetry_payloads = []\n"
        "for box in results[0].boxes:\n"
        "    telemetry_payloads.append({\n"
        "        'class': model.names[int(box.cls)],\n"
        "        'confidence': round(float(box.conf), 2),\n"
        "        'bounding_area': int((box.xyxy[0][2]-box.xyxy[0][0]) * (box.xyxy[0][3]-box.xyxy[0][1]))\n"
        "    })\n\n"
        "del frame # Explicitly destroy raw image data from memory\n"
        "client.publish('smart-shelf/telemetry', json.dumps(telemetry_payloads))"
    )
    code_block_2.style = 'Intense Quote'

    doc.add_heading('2.2 Three-Phase Go-To-Market (GTM) Plan', level=2)
    doc.add_paragraph(
        "• Phase 1 (Alpha): Direct-to-Consumer via Smart Appliance OEMs. Integrate the edge-AI directly into premium refrigerator cameras.\n"
        "• Phase 2 (Beta): Standalone IoT Retrofit Kit. A battery-powered, magnetically mounted wide-angle lens for existing pantries, targeting mid-income households.\n"
        "• Phase 3 (Scale): Enterprise Aggregation. Anonymized decay analytics sold to grocery chains to predict localized food demand and optimize their supply chains, creating a secondary revenue stream."
    )

    # Deliverable 3
    doc.add_heading('Deliverable 3: Optimization Strategy for Metric 03 (Hardware Efficiency)', level=1)
    
    doc.add_heading('3.1 Intel OpenVINO Hardware Optimization', level=2)
    doc.add_paragraph(
        "Running raw PyTorch (.pt) models on edge devices consumes excessive power. To secure hardware efficiency points, "
        "we compile the YOLOv8 model to Intel's OpenVINO IR format, accelerating inference on Intel CPUs and iGPUs while slashing power draw."
    )
    code_block_3 = doc.add_paragraph()
    code_block_3.add_run(
        "# 1. Export YOLOv8 to OpenVINO format\n"
        "from ultralytics import YOLO\n"
        "model = YOLO('yolov8n.pt')\n"
        "model.export(format='openvino', half=True) # FP16 optimization\n\n"
        "# 2. Load the optimized model in production\n"
        "optimized_model = YOLO('yolov8n_openvino_model/')\n"
        "results = optimized_model(frame, device='cpu') # Routes to OpenVINO backend"
    )
    code_block_3.style = 'Intense Quote'

    # Save
    doc.save('C:/Smart-Shelf/Smart_Shelf_Competition_Synopsis.docx')
    print("Successfully generated C:/Smart-Shelf/Smart_Shelf_Competition_Synopsis.docx")

if __name__ == '__main__':
    create_document()
